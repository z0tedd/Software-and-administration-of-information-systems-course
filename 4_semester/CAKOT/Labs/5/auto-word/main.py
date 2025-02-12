from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()


# Add a title page
def add_title_page(doc):
    title = doc.add_heading("Министерство науки и высшего образования РФ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Федеральное государственное бюджетное образовательное учреждение\n"
        "высшего образования\n"
        "«Курский государственный университет»",
        style="Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Кафедра программного обеспечения и администрирования информационных систем\n"
        "Направление подготовки математическое обеспечение и администрирование информационных систем\n"
        "Форма обучения очная",
        style="Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Отчет\nпо лабораторной работе №3\n«Сравнительный анализ алгоритма поиска»",
        style="Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Выполнил:\nстудент группы 213 Файтельсон А.А.\n\n"
        "Проверил:\nассистент кафедры ПОиАИС Овсянников А.В.",
        style="Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Курск, 2024", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )


add_title_page(doc)

# Add sections
doc.add_heading("Цель работы", level=1)
doc.add_paragraph(
    "Изучение алгоритмов поиска элемента в массиве и закрепление навыков в проведении сравнительного анализа алгоритмов."
)

doc.add_heading("Листинг программы", level=1)
code = """package main
import (
    "errors"
    "fmt"
    "math/rand"
    "slices"
    "testing"
    "time"
)
func generateRandomArray(N int) []int {
    rand.Seed(time.Now().UnixNano())
    array := make([]int, N)
    for i := 0; i < N; i++ {
        array[i] = rand.Intn(10000000000)
    }
    return array
}

func NaiveSearch(v int, slice []int) (int, error) {
    for i, value := range slice {
        if value == v {
            return i, nil
        }
    }
    return 0, errors.New("element not found")
}

func FastNaiveSearch(v int, slice []int) (int, error) {
    n := len(slice)
    slice = append(slice, v)
    i := 0
    for slice[i] != v {
        i++
    }
    if i == n {
        return 0, errors.New("element not found")
    }
    return i, nil
}

func BlockSearch(v int, slice []int) (int, error) {
    if len(slice) == 0 {
        return -1, errors.New("element not found")
    }
    medium := len(slice) / 2
    if slice[medium] == v {
        return medium, nil
    }
    if slice[medium] > v {
        result, err := BlockSearch(v, slice[:medium])
        if err != nil {
            return -1, err
        }
        return result, nil
    } else {
        result, err := BlockSearch(v, slice[medium+1:])
        if err != nil {
            return -1, err
        }
        return medium + 1 + result, nil
    }
}

func BinarySearch(v int, slice []int) (int, error) {
    l, r := 0, len(slice)
    for l < r {
        mid := (l + r) / 2
        if slice[mid] == v {
            return mid, nil
        }
        if slice[mid] > v {
            r = mid
        } else {
            l = mid + 1
        }
    }
    return -1, errors.New("element not found")
}

func BenchmarkNaiveSearch(b *testing.B) {
    for _, t := range SizeOfCases {
        b.Run(fmt.Sprintf("size of case: %d", t), func(b *testing.B) {
            b.StopTimer()
            value := 355
            arr := generateRandomArray(t)
            b.StartTimer()
            for i := 0; i < b.N; i++ {
                NaiveSearch(value, arr)
            }
        })
    }
}
"""

# Add the code block with monospaced font
code_paragraph = doc.add_paragraph(style="Normal")
run = code_paragraph.add_run(code)
font = run.font
font.name = "Courier New"  # Monospaced font for code
font.size = Pt(10)

# Add results section
doc.add_heading("Результаты сравнительного анализа алгоритмов", level=1)
doc.add_paragraph("(в микросекундах)")
doc.add_paragraph("Таблица 1 - Максимальное количество операций сравнения")
doc.add_paragraph("Таблица 2 - Среднее количество операций сравнения")

# Add conclusions
doc.add_heading("Вывод по работе", level=1)
doc.add_paragraph(
    "Для решения задачи поиска элемента в отсортированном массиве нужно использовать бинарный поиск, "
    "а в остальных случаях обычный поиск работает приемлемо."
)

# Save the document
doc.save("Отчет_по_лабораторной_работе_3.docx")
